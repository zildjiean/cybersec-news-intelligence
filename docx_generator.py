"""
DOCX Generator – professional Word document using python-docx.
"""
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    RGBColor = None  # guard — colours created lazily inside generate_docx

# Plain tuples (R,G,B) — safe to define at module level without python-docx
SEV_COLORS = {
    "Critical": (0xDC, 0x26, 0x26),
    "High":     (0xEA, 0x58, 0x0C),
    "Medium":   (0xCA, 0x8A, 0x04),
    "Low":      (0x25, 0x63, 0xEB),
    "Info":     (0x6B, 0x72, 0x80),
}
TLP_COLORS = {
    "TLP:CLEAR":        (0x47, 0x55, 0x69),
    "TLP:GREEN":        (0x16, 0xA3, 0x4A),
    "TLP:AMBER":        (0xD9, 0x77, 0x06),
    "TLP:AMBER+STRICT": (0xEA, 0x58, 0x0C),
    "TLP:RED":          (0xDC, 0x26, 0x26),
}

def _rgb(t): return RGBColor(t[0], t[1], t[2])


def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _table_borders(table, color='E4E4E7'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _hrule(doc, color='E2E8F0', sz='6'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _section_heading(doc, text: str, color: RGBColor):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = color
    # Underline via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    # RGBColor stores as 3-byte tuple
    r, g, b = color[0], color[1], color[2]
    bot.set(qn('w:color'), f'{r:02X}{g:02X}{b:02X}')
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_content(doc, text: str):
    if not text:
        p = doc.add_paragraph("ไม่มีข้อมูล")
        if p.runs:
            p.runs[0].font.color.rgb = _rgb((0x94, 0xA3, 0xB8))
        return
    lines = text.replace("•", "\n•").split("\n")
    has_bullets = any(
        l.strip().startswith("•") or
        (len(l.strip()) > 1 and l.strip()[0].isdigit() and l.strip()[1] in '.)' )
        for l in lines
    )
    if has_bullets:
        for line in lines:
            line = line.strip().lstrip("•0123456789.-) ").strip()
            if line:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line)
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(3)
    else:
        for para in text.split("\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                if p.runs:
                    p.runs[0].font.size = Pt(11)
                p.paragraph_format.space_after = Pt(5)


def generate_docx(article: dict, output_path: str) -> bool:
    if not DOCX_AVAILABLE:
        print("[DOCX] python-docx not installed. Run: pip install python-docx")
        return False
    try:
        doc = Document()

        # ── Page setup ──
        for sec in doc.sections:
            sec.page_width  = Cm(21)
            sec.page_height = Cm(29.7)
            sec.left_margin = sec.right_margin = Cm(2.5)
            sec.top_margin  = sec.bottom_margin = Cm(2)

        severity  = article.get("severity", "Info")
        tlp_key   = article.get("tlp") or "TLP:CLEAR"
        category  = article.get("category", "ทั่วไป")
        thai_title   = article.get("thai_title", "ไม่มีชื่อ")
        source_name  = article.get("source_name", "ไม่ทราบ")
        operator     = (article.get("operator") or "").strip()
        url          = article.get("url", "")
        thai_summary = article.get("thai_summary", "")
        thai_content = article.get("thai_content", "")
        thai_impact  = article.get("thai_impact", "")
        thai_rec     = article.get("thai_recommendation", "")
        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")

        sev_rgb = SEV_COLORS.get(severity, SEV_COLORS["Info"])
        tlp_rgb = TLP_COLORS.get(tlp_key, TLP_COLORS["TLP:CLEAR"])

        # ══ HEADER ══
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = h.add_run("CYBERSEC NEWS INTELLIGENCE")
        hr.font.size = Pt(20); hr.font.bold = True
        hr.font.color.rgb = _rgb((0x0A, 0x0E, 0x1A))

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run("รายงานข่าวความมั่นคงปลอดภัยไซเบอร์  |  สำหรับผู้บริหาร, CISO และทีม Security")
        sr.font.size = Pt(9)
        sr.font.color.rgb = _rgb((0x64, 0x74, 0x8B))
        sub.paragraph_format.space_after = Pt(4)

        _hrule(doc, '0A0E1A', '12')

        # ══ CLASSIFICATION BADGES ══
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(8)
        bp.paragraph_format.space_after = Pt(4)

        sr2 = bp.add_run(f"[ {severity.upper()} ]")
        sr2.font.size = Pt(11); sr2.font.bold = True; sr2.font.color.rgb = _rgb(sev_rgb)

        bp.add_run("   ").font.size = Pt(11)

        tr = bp.add_run(f"[ {tlp_key} ]")
        tr.font.size = Pt(11); tr.font.bold = True; tr.font.color.rgb = _rgb(tlp_rgb)

        bp.add_run("   ").font.size = Pt(11)
        cr2 = bp.add_run(category)
        cr2.font.size = Pt(11); cr2.font.color.rgb = _rgb((0x7C, 0x3A, 0xED))

        # ══ ARTICLE TITLE ══
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(8)
        tp.paragraph_format.space_after = Pt(10)
        tr2 = tp.add_run(thai_title)
        tr2.font.size = Pt(16); tr2.font.bold = True
        tr2.font.color.rgb = _rgb((0x0A, 0x0E, 0x1A))

        # ══ METADATA TABLE ══
        meta_rows = [
            ("แหล่งที่มา",          source_name,  None),
            ("ระดับความรุนแรง",     severity,      sev_rgb),
            ("TLP Classification",  tlp_key,       tlp_rgb),
            ("หมวดหมู่",            category,      None),
        ]
        if operator:
            meta_rows.append(("ผู้ดำเนินการ", operator, None))
        meta_rows.append(("วันที่จัดทำ", generated_at, None))

        tbl = doc.add_table(rows=len(meta_rows), cols=2)
        _table_borders(tbl)
        tbl.autofit = False
        tbl.columns[0].width = Cm(4.5)
        tbl.columns[1].width = Cm(11.5)

        for i, (key, val, val_color) in enumerate(meta_rows):
            row = tbl.rows[i]
            kc, vc = row.cells[0], row.cells[1]
            _cell_bg(kc, 'F1F5F9')

            kr = kc.paragraphs[0].add_run(key)
            kr.font.size = Pt(10); kr.font.bold = True
            kr.font.color.rgb = _rgb((0x33, 0x41, 0x55))

            vr = vc.paragraphs[0].add_run(val or "—")
            vr.font.size = Pt(10)
            if val_color:
                vr.font.color.rgb = _rgb(val_color); vr.font.bold = True
            else:
                vr.font.color.rgb = _rgb((0x1E, 0x29, 0x3B))

        # ── Reference URL ──
        doc.add_paragraph()
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(12)
        rl = rp.add_run("Reference URL:  ")
        rl.font.size = Pt(9); rl.font.bold = True
        rl.font.color.rgb = _rgb((0x33, 0x41, 0x55))
        ru = rp.add_run(url or "—")
        ru.font.size = Pt(9)
        ru.font.color.rgb = _rgb((0x0E, 0xA5, 0xE9))

        # ══ CONTENT SECTIONS ══
        sections = [
            ("1. สรุปสำหรับผู้บริหาร (Executive Summary)",
             thai_summary, _rgb((0x0E, 0xA5, 0xE9))),
            ("2. เนื้อหาฉบับเต็ม (ภาษาไทย)",
             thai_content, _rgb((0x7C, 0x3A, 0xED))),
            ("3. การวิเคราะห์ผลกระทบต่อองค์กร",
             thai_impact, _rgb((0xEA, 0x58, 0x0C))),
            ("4. คำแนะนำในการรับมือและป้องกัน",
             thai_rec, _rgb((0x16, 0xA3, 0x4A))),
        ]
        for sec_title, sec_content, sec_color in sections:
            _section_heading(doc, sec_title, sec_color)
            _add_content(doc, sec_content)

        # ══ FOOTER ══
        doc.add_paragraph()
        _hrule(doc, 'CBD5E1', '4')

        conf = doc.add_paragraph()
        conf.paragraph_format.space_before = Pt(4)
        cr3 = conf.add_run("CONFIDENTIAL")
        cr3.font.bold = True; cr3.font.size = Pt(9)
        cr3.font.color.rgb = _rgb((0xDC, 0x26, 0x26))
        cr3b = conf.add_run(" – สำหรับใช้ภายในองค์กรเท่านั้น")
        cr3b.font.size = Pt(9)
        cr3b.font.color.rgb = _rgb((0x64, 0x74, 0x8B))

        foot_text = "จัดทำโดยระบบ CyberSec News Intelligence"
        if operator:
            foot_text += f"  |  ผู้ดำเนินการ: {operator}"
        foot_text += f"  |  วันที่: {generated_at}"
        fp = doc.add_paragraph()
        fr = fp.add_run(foot_text)
        fr.font.size = Pt(8)
        fr.font.color.rgb = _rgb((0x94, 0xA3, 0xB8))

        os.makedirs(
            os.path.dirname(output_path) if os.path.dirname(output_path) else ".",
            exist_ok=True
        )
        doc.save(output_path)
        print(f"[DOCX] Generated: {output_path}")
        return True

    except Exception as e:
        print(f"[DOCX] Failed: {e}")
        import traceback; traceback.print_exc()
        return False
