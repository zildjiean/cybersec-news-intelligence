"""
DOCX Generator – XFINIT v1 Template
Fills the XFINIT v1 .docx template with article data using python-docx + lxml XML manipulation.
"""
import os
import re
import copy
import shutil
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'assets', 'xfinit_v1_template.docx')


# ─────────────────────────────────────────────────────────────────────────────
#  Public entry point
# ─────────────────────────────────────────────────────────────────────────────
def generate_docx_xfinit(article_data: dict, output_path: str) -> bool:
    """
    Fill the XFINIT v1 template with *article_data* and write to *output_path*.
    Returns True on success, False otherwise.
    """
    try:
        from docx import Document
    except ImportError:
        return False

    if not os.path.exists(TEMPLATE_PATH):
        return False

    try:
        import tempfile
        # Copy template to a temp file (open from there so the original is
        # never locked), then save to the desired output_path as a fresh file.
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        try:
            shutil.copy2(TEMPLATE_PATH, tmp_path)
            # Ensure the temp copy is writable (copy2 preserves permissions)
            os.chmod(tmp_path, 0o644)
            doc = Document(tmp_path)
            _fill_template(doc, article_data)
            doc.save(output_path)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
        return True
    except Exception:
        import traceback; traceback.print_exc()
        return False


# ─────────────────────────────────────────────────────────────────────────────
#  Core fill logic
# ─────────────────────────────────────────────────────────────────────────────
def _fill_template(doc, data: dict):
    paras = doc.paragraphs

    title           = data.get('thai_title', '')
    source_name     = data.get('source_name', '')
    url             = data.get('url', '')
    summary         = data.get('thai_summary', '')
    content         = data.get('thai_content', '')
    impact          = data.get('thai_impact', '')
    recommendations = data.get('thai_recommendation', '')

    # ── Find landmark indices ───────────────────────────────────────────────
    def find_idx(text):
        for i, p in enumerate(paras):
            if text in p.text:
                return i
        return None

    sh = find_idx('สรุปสำหรับผู้บริหาร')
    ch = find_idx('เนื้อหาฉบับเต็ม')
    ih = find_idx('การวิเคราะห์ผลกระทบ')
    rh = find_idx('คำแนะนำในการรับมือ')

    if None in (sh, ch, ih, rh):
        raise ValueError(f"Template landmark not found: sh={sh} ch={ch} ih={ih} rh={rh}")

    # ── 1. Single-paragraph replacements ───────────────────────────────────
    # Auto-detect title/source/ref positions: they are the first 3 text-bearing
    # paragraphs before the summary section header (skip image-only paragraphs).
    pre_section = [p for p in paras[:sh]
                   if p.text.strip()
                   and p._element.find(f'.//{{{W}}}drawing') is None]
    if len(pre_section) >= 1: _replace_para(pre_section[0], title)
    if len(pre_section) >= 2: _replace_para(pre_section[1], f"แหล่งที่มา: {source_name}")
    if len(pre_section) >= 3: _replace_para(pre_section[2], f"Reference: {url}")

    # ── 2. Summary section: single para between sh+1 and ch-1 ──────────────
    # Reference formatting from the summary para (sh+1)
    ref_fmt = paras[sh + 1] if sh + 1 < ch else paras[sh]
    _replace_section(paras, sh + 1, ch - 1, [summary], bullet=False, ref_para=ref_fmt)

    # Re-fetch paragraphs after DOM mutation
    paras = doc.paragraphs
    sh = find_idx_list(paras, 'สรุปสำหรับผู้บริหาร')
    ch = find_idx_list(paras, 'เนื้อหาฉบับเต็ม')
    ih = find_idx_list(paras, 'การวิเคราะห์ผลกระทบ')
    rh = find_idx_list(paras, 'คำแนะนำในการรับมือ')

    # ── 3. Full content section ─────────────────────────────────────────────
    ref_fmt = paras[ch + 1] if ch + 1 < ih else paras[ch]
    content_lines = _split_content(content)
    _replace_section(paras, ch + 1, ih - 1, content_lines, bullet=False, ref_para=ref_fmt)

    # Re-fetch
    paras = doc.paragraphs
    ih = find_idx_list(paras, 'การวิเคราะห์ผลกระทบ')
    rh = find_idx_list(paras, 'คำแนะนำในการรับมือ')

    # ── 4. Impact section ───────────────────────────────────────────────────
    ref_fmt = paras[ih + 1] if ih + 1 < rh else paras[ih]
    impact_lines = _split_bullets(impact)
    _replace_section(paras, ih + 1, rh - 1, impact_lines, bullet=True, ref_para=ref_fmt)

    # Re-fetch
    paras = doc.paragraphs
    rh = find_idx_list(paras, 'คำแนะนำในการรับมือ')

    # ── 5. Recommendations section (to end of body) ─────────────────────────
    rec_lines = _split_bullets(recommendations)
    last = len(paras) - 1
    ref_fmt = paras[rh + 1] if rh + 1 <= last else paras[rh]
    _replace_section(paras, rh + 1, last, rec_lines, bullet=True, ref_para=ref_fmt)


def find_idx_list(paras, text):
    for i, p in enumerate(paras):
        if text in p.text:
            return i
    return None


# ─────────────────────────────────────────────────────────────────────────────
#  Paragraph manipulation
# ─────────────────────────────────────────────────────────────────────────────

def _get_rpr(para):
    """Return a deep-clone of the first run's rPr element, or None."""
    for run in para.runs:
        rpr = run._element.find(f'{{{W}}}rPr')
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _get_ppr(para):
    """Return a deep-clone of the paragraph's pPr element, or None."""
    ppr = para._element.find(f'{{{W}}}pPr')
    if ppr is not None:
        return copy.deepcopy(ppr)
    return None


def _replace_para(para, text: str):
    """
    Replace all runs in *para* with a single run containing *text*,
    preserving the first run's character formatting (rPr).
    """
    rpr = _get_rpr(para)
    elem = para._element
    # Remove inline content children
    for child in list(elem):
        local = etree.QName(child.tag).localname
        if local in ('r', 'hyperlink', 'ins', 'del', 'bookmarkStart', 'bookmarkEnd'):
            elem.remove(child)
    # Add new run
    r_elem = etree.SubElement(elem, f'{{{W}}}r')
    if rpr is not None:
        r_elem.insert(0, rpr)
    t_elem = etree.SubElement(r_elem, f'{{{W}}}t')
    t_elem.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _build_para_elem(text: str, ppr, rpr, bullet: bool) -> etree.Element:
    """Build a new <w:p> element with given formatting and text."""
    new_p = etree.Element(f'{{{W}}}p')
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    r_elem = etree.SubElement(new_p, f'{{{W}}}r')
    if rpr is not None:
        r_elem.append(copy.deepcopy(rpr))
    t_elem = etree.SubElement(r_elem, f'{{{W}}}t')
    display = ('• ' + text.lstrip('•- ').strip()) if bullet else text
    t_elem.text = display
    if display and (display[0] == ' ' or display[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p


def _replace_section(paras: list, start: int, end: int,
                     lines: list, bullet: bool, ref_para=None):
    """
    Replace paragraphs paras[start..end] (inclusive) with new paragraphs
    built from *lines*.  Formatting is cloned from *ref_para* (or paras[start]).
    """
    if start > end or start >= len(paras):
        return

    ref = ref_para or paras[start]
    rpr = _get_rpr(ref)
    ppr = _get_ppr(ref)

    # Determine anchor: the element that will come right AFTER our new paragraphs
    body = paras[start]._element.getparent()
    if end + 1 < len(paras):
        anchor = paras[end + 1]._element
    else:
        # Use sectPr or last body child
        anchor = body.find(f'{{{W}}}sectPr')
        if anchor is None:
            anchor = list(body)[-1]

    # Remove existing paragraphs in [start..end]
    for i in range(end, start - 1, -1):
        if i < len(paras):
            elem = paras[i]._element
            if elem.getparent() is not None:
                elem.getparent().remove(elem)

    # Build and insert new paragraph elements (before anchor, in forward order)
    filtered = [l for l in lines if l and l.strip()]
    if not filtered:
        filtered = ['']   # keep at least one empty para to avoid collapsed sections

    for line in filtered:
        new_p = _build_para_elem(line, ppr, rpr, bullet)
        anchor.addprevious(new_p)


# ─────────────────────────────────────────────────────────────────────────────
#  Text parsing helpers
# ─────────────────────────────────────────────────────────────────────────────

def _split_content(text: str) -> list:
    """Split multi-paragraph text into non-empty line list."""
    if not text:
        return ['']
    return [l.strip() for l in text.replace('\r\n', '\n').split('\n') if l.strip()] or ['']


def _split_bullets(text: str) -> list:
    """
    Split bullet-style text into individual items (marker stripped).
    Handles '• ', '- ', '* ', numbered '1. ', and plain paragraphs.
    Each non-empty line becomes its own bullet item.
    """
    if not text:
        return ['']

    lines = text.replace('\r\n', '\n').split('\n')
    items = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Strip common bullet markers if present
        cleaned = re.sub(r'^[•\-\*\d]+[\.\)]*\s*', '', stripped).strip()
        items.append(cleaned if cleaned else stripped)

    return items or ['']
